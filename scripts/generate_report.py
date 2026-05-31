from __future__ import annotations

from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "report_outputs"
ASSET_DIR = OUT_DIR / "report1_assets"
OUT_DOCX = OUT_DIR / "Report1_C_Dashboard_UX_ThreatRadar_v2.docx"


BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F3F5F7"
MID_GRAY = "666666"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Microsoft JhengHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    r.font.size = Pt(9.5)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], LIGHT_BLUE)
        set_cell_text(hdr[i], h, bold=True, color="0B2545")
        if widths:
            hdr[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text)
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()
    return table


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft JhengHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(5)

    for name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 11.5, "1F3A5F"),
    ]:
        s = styles[name]
        s.font.name = "Microsoft JhengHei"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(6)
        s.paragraph_format.space_after = Pt(4)


def add_page_number(section) -> None:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.style = "Footer"
    run = footer.add_run("Report 1 | Page ")
    run.font.name = "Microsoft JhengHei"
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def add_title(doc: Document, title: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(title)
    r.font.name = "Microsoft JhengHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE)
    if subtitle:
        p2 = doc.add_paragraph()
        r2 = p2.add_run(subtitle)
        r2.font.name = "Microsoft JhengHei"
        r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        r2.font.size = Pt(11)
        r2.font.color.rgb = RGBColor.from_string(MID_GRAY)


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT_GRAY) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(title + "\n")
    r.bold = True
    r.font.name = "Microsoft JhengHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(BLUE)
    r2 = p.add_run(body)
    r2.font.name = "Microsoft JhengHei"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    r2.font.size = Pt(9.8)
    doc.add_paragraph()


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Microsoft JhengHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(10)


def para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = "Microsoft JhengHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    r.font.size = Pt(10.5)


def code_image(title: str, path: str, start: int, end: int, out_name: str) -> Path:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    src = (ROOT / path).read_text(encoding="utf-8").splitlines()
    lines = src[start - 1 : end]
    font_paths = [
        r"C:\Windows\Fonts\msjh.ttc",
        r"C:\Windows\Fonts\mingliu.ttc",
        r"C:\Windows\Fonts\consola.ttf",
        r"C:\Windows\Fonts\CascadiaMono.ttf",
        r"C:\Windows\Fonts\cour.ttf",
    ]
    font_path = next((p for p in font_paths if Path(p).exists()), None)
    font = ImageFont.truetype(font_path, 17) if font_path else ImageFont.load_default()
    title_font = ImageFont.truetype(font_path, 16) if font_path else ImageFont.load_default()
    line_h = 24
    width = 1280
    height = 56 + line_h * len(lines) + 24
    img = Image.new("RGB", (width, height), "#1e1e1e")
    draw = ImageDraw.Draw(img)
    draw.rectangle([0, 0, width, 40], fill="#2d2d30")
    draw.text((18, 11), f"{title}  —  {path}:{start}-{end}", fill="#d4d4d4", font=title_font)
    y = 54
    for i, line in enumerate(lines, start):
        draw.text((18, y), f"{i:>4}", fill="#858585", font=font)
        wrapped = wrap(line.expandtabs(4), width=105, replace_whitespace=False) or [""]
        draw.text((78, y), wrapped[0], fill="#dcdcdc", font=font)
        y += line_h
        for cont in wrapped[1:2]:
            draw.text((78, y), cont, fill="#dcdcdc", font=font)
            y += line_h
    out = ASSET_DIR / out_name
    img.save(out)
    return out


def architecture_image() -> Path:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1650, 780), "white")
    draw = ImageDraw.Draw(img)
    font_paths = [r"C:\Windows\Fonts\msjh.ttc", r"C:\Windows\Fonts\arial.ttf"]
    fp = next((p for p in font_paths if Path(p).exists()), None)
    font = ImageFont.truetype(fp, 26) if fp else ImageFont.load_default()
    small = ImageFont.truetype(fp, 19) if fp else ImageFont.load_default()
    caption_font = ImageFont.truetype(fp, 21) if fp else ImageFont.load_default()

    def box(x1, y1, x2, y2, title, subtitle, fill="#eef5fb"):
        draw.rounded_rectangle([x1, y1, x2, y2], radius=14, fill=fill, outline="#1f4e79", width=3)
        draw.text((x1 + 18, y1 + 18), title, fill="#0b2545", font=font)
        draw.text((x1 + 18, y1 + 53), subtitle, fill="#23445f", font=small)

    boxes = {
        "rss": (70, 110, 300, 210, "RSS Feeds", "CISA / iThome / News"),
        "nvd": (70, 330, 300, 430, "NVD API", "Recent CVEs + CVSS"),
        "news_raw": (410, 230, 650, 340, "SQLite news", "raw rows + URL de-dupe"),
        "llm": (750, 230, 990, 340, "Ollama LLM", "threat JSON"),
        "news_done": (1050, 230, 1290, 340, "Analyzed news", "level / CVE / products"),
        "dashboard": (1360, 165, 1600, 275, "Dashboard / UX", "C module controls"),
        "scanner": (1050, 465, 1290, 575, "GitHub Scanner", "dependency matching"),
        "scans": (1360, 465, 1600, 575, "github_scans", "scan history"),
    }
    for spec in boxes.values():
        box(*spec)

    def arrow(start, end, elbow=None, color="#333333"):
        points = [start]
        if elbow:
            if isinstance(elbow[0], tuple):
                points.extend(elbow)
            else:
                points.append(elbow)
        points.append(end)
        for a, b in zip(points, points[1:]):
            draw.line([a, b], fill=color, width=4)
        sx, sy = points[-2]
        ex, ey = end
        if abs(ex - sx) >= abs(ey - sy):
            if ex >= sx:
                head = [(ex, ey), (ex - 14, ey - 8), (ex - 14, ey + 8)]
            else:
                head = [(ex, ey), (ex + 14, ey - 8), (ex + 14, ey + 8)]
        else:
            if ey >= sy:
                head = [(ex, ey), (ex - 8, ey - 14), (ex + 8, ey - 14)]
            else:
                head = [(ex, ey), (ex - 8, ey + 14), (ex + 8, ey + 14)]
        draw.polygon(head, fill=color)

    # Main ingestion and LLM path
    arrow((300, 160), (410, 265), elbow=[(355, 160), (355, 265)])
    arrow((300, 380), (410, 305), elbow=[(355, 380), (355, 305)])
    arrow((650, 285), (750, 285))
    arrow((990, 285), (1050, 285))

    # Dashboard reads analyzed news; scanner also reads analyzed products.
    arrow((1290, 270), (1360, 220))
    arrow((1170, 340), (1170, 465))

    # Scanner stores history; Dashboard reads scan history.
    arrow((1290, 520), (1360, 520))
    arrow((1480, 465), (1480, 275))

    # Labels
    draw.text((430, 185), "insert_news()", fill="#6a4a00", font=small)
    draw.text((760, 185), "analyze_pending_news()", fill="#6a4a00", font=small)
    draw.text((70, 742), "核心資料流：crawler 寫入 news → Ollama 更新為 analyzed news → Dashboard 讀取；GitHub Scanner 讀 analyzed news 並將結果保存到 github_scans。", fill="#333333", font=caption_font)
    out = ASSET_DIR / "architecture_pipeline.png"
    img.save(out)
    return out


def add_picture(doc: Document, image_path: Path, caption: str, width: float = 6.4) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.font.name = "Microsoft JhengHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(MID_GRAY)


def build() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_page_number(doc.sections[0])

    arch = architecture_image()
    img_dashboard = code_image("Dashboard UI controls and event wiring", "dashboard/app.py", 510, 649, "dashboard_controls.png")
    img_db = code_image("Database query layer for C module", "database/db.py", 157, 199, "db_recent_news.png")
    img_status = code_image("Enhanced stats and crawl run record", "database/db.py", 236, 330, "db_stats_crawl_runs.png")
    img_pipeline = code_image("Lock-protected crawl and re-analysis pipeline", "pipeline.py", 25, 73, "pipeline_reanalyze.png")
    img_nvd = code_image("NVD CVSS extraction", "crawler/nvd.py", 61, 84, "nvd_cvss.png")

    # Page 1
    add_title(doc, "Report 1：ThreatRadar Dashboard / UX 模組進度報告", "Applying Large Language Models in Cybersecurity Systems")
    doc.add_paragraph()
    add_table(
        doc,
        ["項目", "內容"],
        [
            ["專題名稱", "ThreatRadar：資安新聞即時威脅雷達"],
            ["負責模組", "C：Dashboard / UX"],
            ["報告類型", "Week 14 Individual Progress Report / Report 1"],
            ["課程要求對應", "個人貢獻追蹤、初步研究、環境建置、模組原型"],
            ["日期", "2026/05/27"],
            ["姓名 / 學號", "（請自行填入）"],
        ],
        widths=[1.5, 4.8],
    )
    add_callout(
        doc,
        "本階段摘要",
        "本報告整理 ThreatRadar 專題目前的系統流程，並聚焦本人負責的 C 模組：Dashboard / UX。"
        "目前 Dashboard 已從原本單純顯示資料，擴充為可搜尋、可排序、可重跑分析、可追蹤掃描歷史、可觀察系統狀態的操作介面。"
        "這使 PoC 不只是後端 pipeline，而是能讓使用者直接理解並操作威脅情報流程的工作台。",
        fill="EAF3F8",
    )
    doc.add_page_break()

    # Page 2
    doc.add_heading("1. 題目要求與個人負責範圍", level=1)
    para(doc, "本課程期末專題要求團隊建立一個以大型語言模型為核心的資安系統 PoC，並在 Week 14、Week 15 提交個人進度報告，最後在 Week 16 完成整合報告與可展示的系統成果。根據專題規格，作品需要呈現 LLM 在資安領域中的實際應用能力，而不是停留在概念展示。")
    para(doc, "ThreatRadar 的問題定義是：資安新聞與 CVE 公告數量龐大，防守者很難每天完整閱讀並判斷哪些漏洞需要立刻處理。因此系統自動收集 RSS 與 NVD 資料，交給本地 LLM 進行威脅等級、CVE、受影響產品與行動建議抽取，最後透過 Dashboard 呈現給使用者。")
    add_table(
        doc,
        ["團隊軌道", "主要工作", "與本人工作的關係"],
        [
            ["A：LLM / 資料品質", "Prompt、模型比較、標註資料集", "Dashboard 顯示 A 輸出的 threat_level、cve_ids、affected_products、action_summary。"],
            ["B：GitHub Scanner", "依賴檔解析與漏洞比對", "Dashboard 提供 GitHub 掃描入口與掃描歷史頁，讓 B 的結果可視化。"],
            ["C：Dashboard / UX", "搜尋、排序、掃描歷史、重新分析、視覺化", "本人主要負責，將後端結果整合成可操作介面。"],
            ["D：DevOps / 工程品質", "Docker、logging、測試、環境設定", "Dashboard 狀態頁與 Docker/Ollama 設定相互配合。"],
        ],
        widths=[1.5, 2.2, 2.7],
    )
    bullet(doc, "Report 1 目標：說明目前系統已能跑通的資料流，並證明 C 模組已具備可展示原型。")
    bullet(doc, "個人貢獻焦點：把資料庫中的威脅情報與 GitHub 掃描結果轉成可以搜尋、篩選、理解與操作的 Dashboard。")
    doc.add_page_break()

    # Page 3
    doc.add_heading("2. 系統整體流程與 Pipeline", level=1)
    para(doc, "目前程式的唯一入口是 main.py。啟動後會先初始化 SQLite，接著立即執行一次 crawl + analyze，然後啟動 APScheduler 定期重跑，最後 launch Gradio Dashboard。這符合題目要求的 PoC 展示：系統不只是一個靜態頁面，而是具有資料收集、LLM 分析、儲存與互動查詢的完整流程。")
    add_picture(doc, arch, "圖 1：ThreatRadar 目前的主要資料流與模組連接", width=6.3)
    add_table(
        doc,
        ["流程階段", "對應檔案", "目前狀態"],
        [
            ["資料收集", "crawler/rss.py、crawler/nvd.py", "RSS 與 NVD 皆透過 insert_news() 寫入 news 表，URL UNIQUE 負責去重。"],
            ["LLM 分析", "analyzer/llm.py", "呼叫 Ollama /api/chat，使用 JSON format 回傳 threat_level、cve_ids、affected_products、action_summary。"],
            ["資料保存", "database/db.py", "SQLite schema 已包含 news、github_scans、crawl_runs，並保留 migration 方式。"],
            ["操作介面", "dashboard/app.py", "Gradio 四個 tab：威脅雷達、GitHub 掃描、掃描歷史、系統狀態。"],
        ],
        widths=[1.2, 2.0, 3.2],
    )
    doc.add_page_break()

    # Page 4
    doc.add_heading("3. C 模組：Dashboard / UX 實作成果", level=1)
    para(doc, "C 模組的核心目標不是單純把資料列出來，而是讓使用者可以快速回答三個問題：目前最嚴重的威脅是什麼？某個 CVE 或產品是否出現在資料庫中？我的 GitHub repo 是否和近期威脅有關？因此 Dashboard 的 UX 被設計成一個威脅情報工作台。")
    add_table(
        doc,
        ["需求", "目前實作", "證據位置"],
        [
            ["搜尋框 + CVE-ID 搜尋", "全文搜尋 title、cve_ids、action_summary、affected_products，並顯示結果筆數。", "dashboard/app.py build_news_html；database/db.py get_recent_news"],
            ["多排序", "支援威脅等級、發布時間、來源、CVSS 分數。", "dashboard/app.py _SORT_MAP；database/db.py _SORT_CLAUSES"],
            ["掃描歷史", "新增掃描歷史 tab，讀取 github_scans，顯示 repo、掃描時間、依賴數、命中數與最高威脅等級。", "dashboard/app.py build_scan_history_html；database/db.py get_scan_history"],
            ["單篇 re-analyze", "透過 pipeline.reanalyze_one() 立即重跑 Ollama，並受全域 lock 保護。", "dashboard/app.py do_reanalyze；pipeline.py reanalyze_one"],
            ["視覺優化", "Dark mode、responsive CSS、CVE chip、CVSS badge、威脅等級 SVG meter、系統狀態頁。", "dashboard/app.py CSS、render_news_card、get_status_html"],
        ],
        widths=[1.25, 3.25, 1.9],
    )
    add_picture(doc, img_dashboard, "圖 2：Dashboard 建構與四個 Gradio 分頁的程式證據", width=6.4)
    doc.add_page_break()

    # Page 5
    doc.add_heading("4. 搜尋、排序、CVE 與 CVSS 視覺化", level=1)
    para(doc, "威脅雷達分頁的資料不是在前端硬切，而是透過 database/db.py 的 get_recent_news() 統一查詢。這樣做的好處是 Dashboard 與 GitHub Scanner 都能使用同一個排序與查詢邏輯，避免 UI 和資料層出現不一致。")
    add_picture(doc, img_db, "圖 3：get_recent_news() 支援搜尋、CVE-only 與多排序", width=6.4)
    para(doc, "CVSS 分數是本次加強後的重要證據。NVD crawler 會從 metrics 中依序嘗試 cvssMetricV31、cvssMetricV30、cvssMetricV2，抽出 baseScore 後寫入 news.cvss_score。Dashboard 再以 badge 顯示 CVSS 分數，並提供 CVSS 分數排序。")
    add_picture(doc, img_nvd, "圖 4：NVD crawler 擷取 CVSS baseScore 並寫入資料庫", width=6.4)
    add_callout(
        doc,
        "UX 設計理由",
        "一般使用者不一定知道 threat_level 的判斷來源，因此 CVSS badge 和可點擊 CVE chip 能提供更直覺的證據鏈：使用者可從 Dashboard 直接跳到 NVD 原始頁面檢查漏洞資訊。",
    )
    doc.add_page_break()

    # Page 6
    doc.add_heading("5. 掃描歷史、重新分析與系統狀態", level=1)
    para(doc, "GitHub 掃描結果會被寫入 github_scans。C 模組新增掃描歷史頁後，過去掃描過的 repo 不再只是一次性結果，而是可追蹤的歷史資料。每筆紀錄會解析 dependencies 與 matched_cves JSON，計算命中數與最高威脅等級，並用 details 區塊展開命中套件。")
    para(doc, "重新分析功能是 Dashboard 和 LLM pipeline 的重要連接。原本如果 LLM 回答不理想，使用者只能等待下一輪排程。現在可從下拉選單挑選單篇新聞，呼叫 reanalyze_one() 立即重跑。同時，pipeline 使用同一把 module-level lock，避免手動 re-analyze 與排程 crawl 同時打 Ollama。")
    add_picture(doc, img_pipeline, "圖 5：run_crawl_cycle() 與 reanalyze_one() 共用全域 lock", width=6.4)
    para(doc, "系統狀態頁則負責把 PoC 是否健康地顯示出來，包括已分析新聞、待分析、分析失敗、GitHub 掃描記錄、最近爬取時間與 Ollama 模型狀態。這對 demo 很重要，因為它能快速證明系統不是只有靜態資料，而是有執行狀態與可觀測性。")
    add_picture(doc, img_status, "圖 6：get_enhanced_stats() 與 crawl_runs 支援真實最近爬取時間", width=6.4)
    doc.add_page_break()

    # Page 7
    doc.add_heading("6. 安全性、穩定性與程式品質", level=1)
    para(doc, "Dashboard 因為使用 gr.HTML 顯示 RSS、LLM、GitHub repo URL 等外部資料，因此本階段加入 html.escape() 和 _safe_url()，避免外部字串直接變成 HTML。這對資安系統特別重要，因為展示層如果忽略輸出編碼，反而可能引入 XSS 類型的風險。")
    add_table(
        doc,
        ["品質項目", "目前處理方式", "意義"],
        [
            ["避免重複爬取 / 重複 LLM", "pipeline.py 使用 threading.Lock，非阻塞取得。", "避免 scheduler 與手動按鈕同時分析同一批資料。"],
            ["LLM 失敗重試", "mark_analysis_failed() 只在重試達上限後標記分析失敗。", "避免暫時性 Ollama 故障造成資料永久流失。"],
            ["HTML 輸出安全", "Dashboard 對 title、source、action、repo URL 等欄位做 escape / URL 限制。", "降低外部資料污染 UI 的風險。"],
            ["Docker 整合", "docker-compose.yml 以 OLLAMA_BASE_URL=http://ollama:11434 指向服務名稱。", "解決 container 內 localhost 指向錯誤服務的問題。"],
            ["程式檢查", "已執行 python -m compileall -q .，結果 OK。", "確認目前 Python 檔案語法可通過編譯。"],
        ],
        widths=[1.35, 2.75, 2.3],
    )
    add_callout(
        doc,
        "目前限制",
        "GitHub Scanner 目前仍主要支援 requirements.txt 與 package.json，pom.xml、go.mod、Cargo.toml 等格式仍屬 B 模組後續擴充。"
        "另外，LLM 對 affected_products 的命名品質會直接影響 scanner 命中率，因此 A 模組的 prompt/evaluation 仍會影響 C 模組呈現的結果品質。",
        fill="FFF4D6",
    )
    doc.add_page_break()

    # Page 8
    doc.add_heading("7. 下階段計畫與結論", level=1)
    para(doc, "Report 1 階段已完成可操作的 Dashboard 原型，並把後端 pipeline、SQLite 查詢、GitHub scan history、re-analysis 與系統狀態串接起來。此成果對整體專題的價值在於：它讓 LLM 分析結果不只是存在資料庫中，而是能被使用者搜尋、驗證、排序與重新分析。")
    add_table(
        doc,
        ["下階段工作", "預期成果"],
        [
            ["加入實際畫面截圖", "在本報告中的程式碼證據圖之外，補上 VS Code、Gradio Dashboard 與 scan history 的實際截圖。"],
            ["配合 A 模組 fixture", "使用固定測試新聞資料驗證搜尋、CVE chip、CVSS badge、re-analyze 的展示流程。"],
            ["配合 B 模組擴充格式", "當 scanner 支援更多依賴格式後，掃描歷史頁可展示更多 ecosystem 的命中結果。"],
            ["UX polish", "改善行動版版面、空狀態文案、loading 狀態與錯誤訊息，讓 demo 更穩定。"],
            ["報告 2 準備", "整理實測截圖、使用案例、功能前後比較與限制分析。"],
        ],
        widths=[2.0, 4.4],
    )
    add_callout(
        doc,
        "結論",
        "本人負責的 C：Dashboard / UX 已達成 Report 1 的初步原型要求。"
        "目前 Dashboard 不只呈現威脅新聞，也能作為整個 ThreatRadar 的控制台，支援人工查詢、漏洞證據追蹤、掃描結果回看、LLM 重新分析與系統健康檢查。"
        "後續將以實際 demo 截圖與固定測試資料補強報告證據，並配合 A/B/D 模組完成 Week 15 與 Week 16 的整合展示。",
        fill="EAF3F8",
    )

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
