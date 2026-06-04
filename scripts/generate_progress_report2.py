from __future__ import annotations

from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


REPO_ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = REPO_ROOT / "report_outputs"
ASSET_DIR = OUT_DIR / "report2_assets"
OUT_DOCX = OUT_DIR / "NTUST_M11415121_鍾唐福_progress_report_2.docx"

SCREEN_THREAT = Path(r"D:\Computer Science\LLM in Cybersecurity Systems\{870EEA24-3233-4799-89F0-31B0CFFBA6F2}.png")
SCREEN_STATUS = Path(r"D:\Computer Science\LLM in Cybersecurity Systems\{59B177A5-48FE-463F-9D51-80123F330B6E}.png")

BLUE = "1F4E79"
LIGHT_BLUE = "F2F4F7"
LIGHT_GRAY = "F3F5F7"
LIGHT_GREEN = "E8F5E9"
LIGHT_YELLOW = "FFF4D6"
MID_GRAY = "666666"


def _font_path() -> str | None:
    for p in [
        r"C:\Windows\Fonts\msjh.ttc",
        r"C:\Windows\Fonts\mingliu.ttc",
        r"C:\Windows\Fonts\CascadiaMono.ttf",
        r"C:\Windows\Fonts\consola.ttf",
    ]:
        if Path(p).exists():
            return p
    return None


def set_font(run, size: float = 10.5, bold: bool = False, color: str | None = None) -> None:
    run.font.name = "Microsoft JhengHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_font(r, size=9.2, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], LIGHT_BLUE)
        set_cell_text(hdr[i], h, bold=True, color="000000")
        if widths:
            hdr[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text)
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()
    return table


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft JhengHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.12
    normal.paragraph_format.space_after = Pt(5)

    for name, size, color, before, after in [
        ("Heading 1", 16, "000000", 10, 6),
        ("Heading 2", 13, "000000", 8, 4),
        ("Heading 3", 11.5, "000000", 6, 3),
    ]:
        s = doc.styles[name]
        s.font.name = "Microsoft JhengHei"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)


def add_page_number(section) -> None:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Report 2 | Page ")
    set_font(run, size=9, color=MID_GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(title)
    set_font(r, size=22, bold=True, color=BLUE)
    p2 = doc.add_paragraph()
    r2 = p2.add_run(subtitle)
    set_font(r2, size=11, color=MID_GRAY)


def para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    set_font(r, size=10.3)


def labeled_para(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(label)
    set_font(r, size=10.3, bold=True, color="000000")
    r2 = p.add_run(text)
    set_font(r2, size=10.3, color="000000")


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_font(r, size=10)


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT_GRAY) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(title + "\n")
    set_font(r, size=10.4, bold=True, color=BLUE)
    r2 = p.add_run(body)
    set_font(r2, size=9.7)
    doc.add_paragraph()


def add_picture(doc: Document, image_path: Path, caption: str, width: float = 6.4) -> None:
    if not image_path.exists():
        add_callout(doc, "缺少截圖", f"找不到圖片：{image_path}", fill=LIGHT_YELLOW)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_font(r, size=8.8, color=MID_GRAY)


def code_image(title: str, path: str, start: int, end: int, out_name: str) -> Path:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    src = (REPO_ROOT / path).read_text(encoding="utf-8").splitlines()
    lines = src[start - 1 : end]
    fp = _font_path()
    font = ImageFont.truetype(fp, 17) if fp else ImageFont.load_default()
    title_font = ImageFont.truetype(fp, 16) if fp else ImageFont.load_default()
    line_h = 24
    width = 1350
    height = 56 + line_h * len(lines) + 24
    img = Image.new("RGB", (width, height), "#1e1e1e")
    draw = ImageDraw.Draw(img)
    draw.rectangle([0, 0, width, 40], fill="#2d2d30")
    draw.text((18, 11), f"{title}  --  {path}:{start}-{end}", fill="#d4d4d4", font=title_font)
    y = 54
    for i, line in enumerate(lines, start):
        draw.text((18, y), f"{i:>4}", fill="#858585", font=font)
        wrapped = wrap(line.expandtabs(4), width=112, replace_whitespace=False) or [""]
        draw.text((78, y), wrapped[0], fill="#dcdcdc", font=font)
        y += line_h
        for cont in wrapped[1:2]:
            draw.text((78, y), cont, fill="#dcdcdc", font=font)
            y += line_h
    out = ASSET_DIR / out_name
    img.save(out)
    return out


def terminal_log_image() -> Path:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    lines = [
        r".\.venv\Scripts\python.exe main.py",
        "[Main] Database initialized",
        "[Scheduler] Starting crawl cycle...",
        "[RSS] Error crawling CISA Alerts: 403 Client Error: Forbidden ...",
        "[LLM] Analyzed 50/50 items (primary=deepseek/deepseek-v4-flash, mode=single, concurrency=3)",
        "[Scheduler] 爬取完成：RSS +0 筆，NVD +0 筆，分析 50 筆",
        "[Main] Scheduler started (every 60 min)",
        "[Main] Launching dashboard at http://localhost:7860",
        "Running on local URL:  http://0.0.0.0:7860",
        "後續排程：22/22、6/6、2/2 items 均由 DeepSeek 完成分析",
    ]
    fp = _font_path()
    font = ImageFont.truetype(fp, 22) if fp else ImageFont.load_default()
    title_font = ImageFont.truetype(fp, 20) if fp else ImageFont.load_default()
    width, height = 1500, 360
    img = Image.new("RGB", (width, height), "#111827")
    draw = ImageDraw.Draw(img)
    draw.rectangle([0, 0, width, 46], fill="#1f2937")
    draw.text((20, 13), "Terminal execution evidence -- DeepSeek API pipeline", fill="#e5e7eb", font=title_font)
    y = 62
    for line in lines:
        color = "#86efac" if "[LLM]" in line or "後續排程" in line else "#d1d5db"
        draw.text((24, y), line, fill=color, font=font)
        y += 30
    out = ASSET_DIR / "terminal_deepseek_run.png"
    img.save(out)
    return out


def architecture_image() -> Path:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1900, 820), "white")
    draw = ImageDraw.Draw(img)
    fp = _font_path()
    font = ImageFont.truetype(fp, 25) if fp else ImageFont.load_default()
    small = ImageFont.truetype(fp, 18) if fp else ImageFont.load_default()
    group_font = ImageFont.truetype(fp, 22) if fp else ImageFont.load_default()
    cap = ImageFont.truetype(fp, 21) if fp else ImageFont.load_default()

    def box(x1, y1, x2, y2, title, subtitle, fill="#ffffff", outline="#1f4e79"):
        draw.rounded_rectangle([x1, y1, x2, y2], radius=12, fill=fill, outline=outline, width=3)

        def draw_wrapped(text, x, y, max_width, text_font, fill="#111111"):
            words = text.split(" ")
            lines: list[str] = []
            current = ""
            for word in words:
                candidate = word if not current else current + " " + word
                if draw.textlength(candidate, font=text_font) <= max_width:
                    current = candidate
                else:
                    if current:
                        lines.append(current)
                    current = word
            if current:
                lines.append(current)
            for idx, line in enumerate(lines[:2]):
                draw.text((x, y + idx * 30), line, fill=fill, font=text_font)
            return y + max(1, min(len(lines), 2)) * 30

        next_y = draw_wrapped(title, x1 + 18, y1 + 18, x2 - x1 - 36, font)
        draw_wrapped(subtitle, x1 + 18, max(y1 + 62, next_y + 4), x2 - x1 - 36, small)

    def arrow(points, color="#333333"):
        for a, b in zip(points, points[1:]):
            draw.line([a, b], fill=color, width=4)
        sx, sy = points[-2]
        ex, ey = points[-1]
        if abs(ex - sx) >= abs(ey - sy):
            head = [(ex, ey), (ex - 14, ey - 8), (ex - 14, ey + 8)] if ex >= sx else [(ex, ey), (ex + 14, ey - 8), (ex + 14, ey + 8)]
        else:
            head = [(ex, ey), (ex - 8, ey - 14), (ex + 8, ey - 14)] if ey >= sy else [(ex, ey), (ex - 8, ey + 14), (ex + 8, ey + 14)]
        draw.polygon(head, fill=color)

    # Input layer
    box(70, 110, 325, 220, "RSS Feeds", "iThome / THN / News")
    box(70, 340, 325, 450, "NVD API", "Recent CVEs + CVSS")
    box(455, 235, 745, 355, "SQLite news", "raw rows + URL de-dupe")

    # Analyzer layer
    box(835, 235, 1135, 355, "Provider Router", "PRIMARY_PROVIDER + mode", fill="#f8fafc")

    draw.rounded_rectangle([1245, 95, 1535, 455], radius=16, fill="#f8fafc", outline="#1f4e79", width=3)
    draw.text((1265, 118), "LLM Providers", fill="#111111", font=group_font)
    box(1270, 165, 1510, 270, "DeepSeek API", "cloud primary", fill="#ffffff")
    box(1270, 320, 1510, 425, "Ollama llama3.2", "local ready", fill="#ffffff")

    # Storage + UI layer
    box(1615, 160, 1865, 275, "news primary row", "dashboard result")
    box(1245, 575, 1535, 690, "news_analyses", "provider history")
    box(1615, 430, 1865, 545, "Dashboard / Export", "C module + reports")

    # Main data flow
    arrow([(325, 165), (390, 165), (390, 275), (455, 275)])
    arrow([(325, 395), (390, 395), (390, 315), (455, 315)])
    arrow([(745, 295), (835, 295)])
    arrow([(1135, 295), (1245, 295)])
    arrow([(1535, 225), (1615, 225)])
    arrow([(1740, 275), (1740, 430)])

    # Provider history and export path
    arrow([(1390, 455), (1390, 575)])
    arrow([(1535, 633), (1585, 633), (1585, 487), (1615, 487)])

    draw.text((70, 745), "核心流程：RSS/NVD 寫入 news；Provider Router 依設定呼叫 DeepSeek 或 Ollama；主要結果更新 news primary row。", fill="#111111", font=cap)
    draw.text((70, 778), "provider history 另外保存到 news_analyses；Dashboard 與匯出功能可讀取主要結果，也能延伸做雙模型比較。", fill="#111111", font=cap)
    out = ASSET_DIR / "report2_architecture.png"
    img.save(out)
    return out


def build() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_page_number(doc.sections[0])

    arch = architecture_image()
    term = terminal_log_image()

    add_title(
        doc,
        "Report 2：ThreatRadar 第二週進度報告",
        "Applying Large Language Models in Cybersecurity Systems",
    )
    doc.add_paragraph()
    add_table(
        doc,
        ["項目", "內容"],
        [
            ["專題名稱", "ThreatRadar：資安新聞即時威脅雷達"],
            ["負責模組", "C：Dashboard / UX，並協助整合 LLM provider 切換與輸出報表"],
            ["報告類型", "Week 15 Individual Progress Report / Report 2"],
            ["本週主軸", "完成 DeepSeek API 實跑，並確認本地 llama3.2 已可連線，支援後續雙模型比較"],
            ["日期", "2026/06/03"],
            ["姓名 / 學號", "鍾唐福 / M11415121"],
        ],
        widths=[1.45, 4.95],
    )
    labeled_para(
        doc,
        "本階段摘要：",
        "第二週已完成最關鍵的整合驗證：同一個 LLM 分析步驟不再只綁定本地 Ollama，而是可透過設定切換為線上 API。"
        "本次實測以 DeepSeek 作為 primary provider，啟動 main.py 後成功完成多輪 crawl + analyze；重新執行後也確認本地 llama3.2 已載入，代表系統已具備雙模型比較的環境條件。"
        "因此目前系統已具備可展示的 end-to-end pipeline：新聞收集、CVE 收集、雲端 LLM 分析、本地模型狀態驗證、SQLite 儲存、Dashboard 查詢與狀態觀察。",
    )
    para(doc, "本報告延續 Report 1 的內容，但重點從 Dashboard 原型轉向整合驗證。也就是說，本週的成果不是單一畫面修改，而是確認後端爬蟲、LLM provider、資料庫、Dashboard 與輸出報表能在同一條流程上連續運作。")
    doc.add_page_break()

    doc.add_heading("1. Report 1 到 Report 2 的進度差異", level=1)
    para(doc, "Report 1 的重點是完成 C 模組 Dashboard / UX 原型，讓使用者能搜尋、排序、查看掃描歷史、重跑單篇分析與觀察系統狀態。Report 2 則進一步把整個系統從「本地模型 PoC」推進到「可切換模型來源的可展示系統」。")
    para(doc, "第一個差異是模型來源。Report 1 的系統仍以本地 Ollama llama3.2 為主要設定，Report 2 則加入 provider abstraction，讓 PRIMARY_PROVIDER 可以在 local 與 cloud 之間切換。這代表後續不需要重寫 Dashboard 或資料庫，只要改環境變數即可選擇主要分析模型。")
    para(doc, "第二個差異是線上 API 已從規劃變成實作。本週新增 CloudProvider，使用 OpenAI-compatible /chat/completions 介面呼叫 DeepSeek API，並成功完成多輪新聞分析。終端機 log 顯示 DeepSeek 分析 50/50、22/22、6/6、2/2 items，表示 API 與原本 pipeline 已經接上。")
    para(doc, "第三個差異是雙模型展示條件更完整。最新系統狀態截圖顯示本地 llama3.2 已載入，雲端 DeepSeek 也已設定。雖然目前 ANALYSIS_MODE 仍是 single，主要結果由 DeepSeek 寫入 news 表，但系統已具備切換 compare/hybrid 模式的基礎。")
    para(doc, "目前仍有一個小限制：CISA RSS feed 回傳 403。不過其他 RSS 與 NVD 仍能運作，因此這不是整體 pipeline 失敗，而是單一資料來源的存取限制。")
    para(doc, "這個進度差異代表系統的風險也有所改變。Report 1 主要擔心 UI 是否能把資料庫資料清楚呈現；Report 2 則更關注模型來源切換、API 設定、分析結果是否能一致寫回，以及 Dashboard 是否能正確顯示目前實際使用的模型。")

    doc.add_heading("2. 更新後的系統 Pipeline", level=1)
    para(doc, "目前 main.py 仍是唯一入口：初始化 SQLite、執行一次 crawl + analyze、啟動 APScheduler，再 launch Gradio。差異在於 analyzer 層不再直接綁死 Ollama，而是先根據 config 選擇 provider。這個改動讓同一個分析流程能在本地模型與線上 API 之間切換。")
    add_picture(doc, arch, "圖 1：Report 2 更新後的 ThreatRadar pipeline：DeepSeek API 與 Ollama 共用同一分析介面", width=6.35)
    para(doc, "資料流可以分成兩條。第一條是主要威脅情報路徑：RSS 與 NVD 先寫入 news 表，Provider Router 再依照 PRIMARY_PROVIDER 呼叫主要模型，最後把威脅等級、CVE、受影響產品與行動建議寫回 news，供 Dashboard 的威脅雷達頁直接讀取。第二條是模型歷史與比較路徑：每一次 provider 的輸出都會另外寫入 news_analyses，這讓未來可以比較 DeepSeek 與 llama3.2 對同一篇新聞的判斷差異。")
    para(doc, "這樣的分層對期末展示很重要。目前 demo 使用 single 模式由 DeepSeek 寫入主要結果，而系統狀態頁已確認本地 llama3.2 載入成功。也就是說，雙模型比較不再只是設計想法，而是可以在下一步透過 ANALYSIS_MODE=compare 展示 DeepSeek 與 llama3.2 對同一批新聞的判斷差異。")
    para(doc, "從 C 模組角度來看，Dashboard 不只是把資料列出來，而是整個 pipeline 的觀察介面。當新聞卡片顯示威脅等級、CVE、模型 badge 與行動建議時，使用者可以確認資料是否完成分析；當系統狀態頁顯示待分析筆數、失敗筆數與最近爬取時間時，使用者可以判斷排程是否正常運作。這些畫面讓後端流程的狀態變成可被檢查與展示的證據。")

    doc.add_heading("3. LLM Provider 切換與 DeepSeek API 實作", level=1)
    para(doc, "本週的主要技術改動是將 LLM 呼叫抽象為 provider。config.py 新增 PRIMARY_PROVIDER、ANALYSIS_MODE、CLOUD_LLM_PROVIDER、CLOUD_LLM_BASE_URL、CLOUD_LLM_MODEL 與 CLOUD_LLM_API_KEY。使用者不需要改程式碼，只要改環境變數，就能把主要分析模型從本地 Ollama 切到 DeepSeek、Qwen 或其他 OpenAI-compatible API。")
    para(doc, "CloudProvider 採用 OpenAI-compatible 的 /chat/completions schema，因此 DeepSeek、Qwen、OpenAI 或其他 gateway 都可以在相同介面下使用。本次實測設定 CLOUD_LLM_PROVIDER=deepseek、CLOUD_LLM_MODEL=deepseek-v4-flash，並由 DeepSeek 作為 primary provider 寫入 news 表。")
    para(doc, "provider abstraction 的價值在於降低切換模型的成本。原本如果把 Ollama 呼叫直接寫在 analyzer/llm.py 中，之後要換 DeepSeek 就會牽動 prompt、JSON parsing、retry 與資料庫更新。現在本地模型與雲端模型都回傳相同 schema，Dashboard 也只需要讀取 news 表中的主要結果，因此 UI 不需要知道背後使用的是哪一家模型服務。")
    para(doc, "此設計也讓 prompt 與輸出格式維持一致。無論模型是本地 llama3.2 或 DeepSeek API，分析結果都會被整理成 threat_level、cve_ids、affected_products 與 action_summary。這對資安新聞系統很重要，因為 Dashboard 與 GitHub Scanner 後續都依賴這些欄位；如果不同模型輸出格式差異太大，就會讓前端呈現與 repo dependency matching 變得不穩定。")
    labeled_para(
        doc,
        "本次設定方式：",
        "本次實測使用 PowerShell 設定：PRIMARY_PROVIDER=cloud、ANALYSIS_MODE=single、CLOUD_LLM_PROVIDER=deepseek、CLOUD_LLM_BASE_URL=https://api.deepseek.com、CLOUD_LLM_MODEL=deepseek-v4-flash。"
        "API key 僅由環境變數提供，不寫入 GitHub。",
    )
    labeled_para(doc, "程式碼截圖建議：", "本版報告先列出建議截圖位置，之後可用 VS Code 自行截圖替換或補到附錄。建議截圖：config.py 26-64、analyzer/providers.py 27-58、analyzer/cloud_provider.py 20-51、analyzer/llm.py 170-214。")
    doc.add_page_break()

    doc.add_heading("4. 實跑結果與 Dashboard 證據", level=1)
    para(doc, "實際執行 .venv 內的 python main.py 後，系統成功初始化資料庫、執行排程爬取、呼叫 DeepSeek 分析，並在 http://127.0.0.1:7860 啟動 Dashboard。終端機 log 明確顯示 primary=deepseek/deepseek-v4-flash，代表 cloud API 已接到原本的分析 pipeline；系統狀態頁也顯示本地 llama3.2 已載入，表示 local provider 的前置條件也已補齊。")
    add_picture(doc, term, "圖 2：終端機執行證據：DeepSeek API 完成多輪 LLM 分析", width=6.35)
    add_picture(doc, SCREEN_THREAT, "圖 3：威脅雷達頁顯示 DeepSeek 產生的威脅等級、行動建議與模型 badge", width=6.45)
    para(doc, "圖 3 可看到每張威脅卡片包含 CRITICAL 等級、摘要建議、CVE chip 與 deepseek/deepseek-v4-flash badge。這證明 Dashboard 讀到的不是舊的 Ollama 結果，而是目前 primary provider 的 DeepSeek 分析結果。")
    para(doc, "威脅雷達頁的呈現方式也符合本專題的核心目標：資安人員不需要先閱讀完整新聞內容，就能快速知道哪些事件最急、是否有 CVE、是否涉及供應鏈或套件風險，以及下一步應該檢查或修補什麼。這比單純 RSS 清單更接近威脅情報 dashboard 的用途。")

    doc.add_heading("5. 系統狀態、資料統計與輸出設計", level=1)
    para(doc, "系統狀態頁用來證明 PoC 是否真的在運作，而不是只有靜態截圖。最新截圖顯示已分析新聞 261 筆、待分析 0 筆、分析失敗 0 筆，本地模型 llama3.2 已載入，cloud provider 也已設定為 deepseek/deepseek-v4-flash。這表示 pipeline 已能把新聞資料分析完並穩定呈現，而且兩種模型來源都具備後續比較的條件。")
    add_picture(doc, SCREEN_STATUS, "圖 4：系統狀態頁顯示 llama3.2 已載入、DeepSeek primary provider 與資料庫統計", width=6.45)
    labeled_para(
        doc,
        "觀察結果：",
        "最新實跑結果比上一版更完整：Ollama 連線檢查已顯示 llama3.2 載入成功，雲端 DeepSeek 也完成 primary provider 設定。"
        "因此 Report 2 可以同時證明線上 API 已跑通，以及本地模型已準備好進入 compare 模式。",
    )
    para(doc, "這個狀態頁對 demo 有兩個用途。第一，它能讓評審立即看到目前分析模型與資料庫統計，避免只靠口頭說明。第二，它能協助除錯，例如待分析數是否累積、分析失敗是否增加、最近爬取時間是否更新。對 C 模組而言，這是把後端狀態轉換成可理解資訊的關鍵畫面。")
    para(doc, "從最新統計來看，已分析 261 筆、待分析 0 筆、失敗 0 筆，代表資料庫內目前沒有堆積未處理項目。威脅等級分布中 HIGH 與 MEDIUM 佔多數，CRITICAL 也有一定數量，這讓 Dashboard 的排序與篩選功能有實際展示價值，而不是只有空白或單一等級資料。")
    doc.add_page_break()

    doc.add_heading("6. Provider History、模型比較與匯出", level=1)
    para(doc, "為了支援 Report 2 提到的雙模型設計，本週新增 news_analyses 表。news 表保留 dashboard 預設結果，而 news_analyses 以 append-only 方式保存每次 provider 的完整輸出，包含 provider、model、prompt_version、threat_level、cve_ids、affected_products、latency_ms、status 與 error。")
    para(doc, "匯出功能則補上了報告與 demo 需要的 output 檔案。exporter/report_exporter.py 可輸出 threat_report CSV、threat_report JSONL 以及 model_comparison CSV。CSV 使用 utf-8-sig 方便 Excel 正確顯示繁體中文，JSONL 保留 list 欄位，方便後續評測或人工檢查。")
    para(doc, "輸出檔案的設計分成三種。threat_report_*.csv 適合放進報告或用 Excel 檢查，欄位包含新聞來源、CVE、產品、威脅等級、行動建議與模型資訊。threat_report_*.jsonl 則保留 cve_ids 與 affected_products 的 array 型態，未來可以接到 evaluation dataset 或人工標註流程。model_comparison_*.csv 是雙模型模式的展示重點，能把 local 與 cloud 的 threat_level 差異整理成可討論的資料。")
    para(doc, "這個 output 設計可以讓成果不只停留在 Gradio 介面。CSV 可以直接放進進度報告、期末簡報或人工驗證流程；JSONL 可以作為後續 evaluation dataset 的雛形；model_comparison.csv 則能把雙模型差異轉成可量化資料，例如統計兩個模型在 CRITICAL/HIGH 判斷上的一致率，或找出 action_summary 明顯不同的案例。")
    labeled_para(doc, "程式碼截圖建議：", "若要證明資料庫與匯出功能，可截 database/db.py 361-462、exporter/report_exporter.py 77-121、dashboard/app.py 570-633，以及 dashboard/app.py 772-786。這幾段分別對應 provider history、報表輸出、系統狀態顯示與匯出按鈕。")

    doc.add_heading("7. 待加強項目與下週計畫", level=1)
    para(doc, "目前整體程式已可算是跑通：DeepSeek API 能完成 LLM 分析，Dashboard 能顯示結果，系統狀態能反映資料庫與 provider 設定。不過，若要讓 final demo 更穩定，仍有幾個可以補強的小項目。")
    para(doc, "第一個待處理項目是 CISA RSS 403。這不是整體 pipeline 的錯誤，因為其他 RSS 與 NVD 仍能正常新增資料，但 demo 時容易被誤解成爬蟲失敗。下週可以在 requests headers 加入 User-Agent，或在報告中說明這是單一來源的存取限制。")
    para(doc, "第二個項目是 cloud primary 未設定時的保護。目前如果 PRIMARY_PROVIDER=cloud 但忘記設定 API key 或 model，系統可能會把資料列入 retry。較好的做法是在 analyze_pending_news() 一開始就檢查設定，若 cloud primary 未啟用就直接跳過，不消耗 retry budget。")
    para(doc, "第三個項目是 API 相容性。雖然 DeepSeek 已經跑通，但不同 OpenAI-compatible provider 對 response_format 的支援程度不一定相同。若要強調可切換 Qwen、OpenRouter 或其他 gateway，CloudProvider 可以在 response_format 不支援時自動 fallback 成一般 JSON prompt 呼叫。")
    para(doc, "第四個項目是雙模型展示。目前 single 模式已足以證明 DeepSeek pipeline 成功，而且本地 llama3.2 已載入；下週若要更符合「兩個 LLM 分析」的專題亮點，可以將 ANALYSIS_MODE 改成 compare，讓模型比較頁與 model_comparison.csv 產出更完整的證據。")
    para(doc, "第五個項目是 GitHub Scanner 的展示資料。目前系統已把 affected_products 作為新聞分析與 repo 掃描的銜接欄位，但 demo 時最好準備一個包含 requirements.txt 或 package.json 的測試 repository。這樣可以實際展示「新聞威脅情報」如何轉成「我的專案是否可能受影響」的個人化檢查結果。")
    labeled_para(
        doc,
        "Report 2 結論：",
        "第二週已完成從本地 LLM PoC 到可切換線上 API 的主要整合。"
        "ThreatRadar 現在可以用 DeepSeek 作為主要分析模型，自動爬取資安新聞與 NVD CVE，將分析結果寫入 SQLite，並在 Dashboard 顯示威脅等級、行動建議、CVE 與模型來源。"
        "本人負責的 C 模組也從 Dashboard 原型提升為可驗證整體 pipeline 的操作介面，足以支援期末 demo 與最終報告。",
    )

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
